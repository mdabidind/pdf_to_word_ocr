import os
import sys
import tempfile
import subprocess
from pathlib import Path
import pdfplumber
import tabula
from docx import Document

# === Tools ===
ROOT = os.path.abspath(os.getcwd())
TOOLS = os.path.join(ROOT, "tools")

TESSERACT = os.environ.get("TESSERACT_PATH") or os.path.join(TOOLS, "tesseract", "tesseract.exe")
POPPLER_BIN = os.environ.get("POPPLER_PATH") or os.path.join(TOOLS, "poppler", "bin")
JAVA_EXE = os.environ.get("JAVA_PATH") or os.path.join(TOOLS, "java", "bin", "java.exe")
TABULA_JAR = os.environ.get("TABULA_JAR") or os.path.join(TOOLS, "tabula", "tabula.jar")
TESSDATA_PREFIX = os.environ.get("TESSDATA_PREFIX") or os.path.join(TOOLS, "tesseract", "tessdata")


# --- Helper: detect if PDF has selectable text ---
def _pdf_has_text(pdf_path, sample_pages=3):
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i in range(min(sample_pages, len(pdf.pages))):
                if (pdf.pages[i].extract_text() or "").strip():
                    return True
    except Exception:
        pass
    return False


# --- Run OCRmyPDF ---
def _ocr_to_searchable_pdf(input_pdf, output_pdf, lang="eng"):
    env = os.environ.copy()
    if os.path.exists(TESSERACT):
        env["TESSERACT_PATH"] = TESSERACT
    if os.path.exists(TESSDATA_PREFIX):
        env["TESSDATA_PREFIX"] = TESSDATA_PREFIX

    cmd = [
        sys.executable, "-m", "ocrmypdf",
        "--skip-text", "--rotate-pages", "--deskew", "--clean",
        "--remove-background", "--optimize", "1",
        "--language", lang,
        input_pdf, output_pdf
    ]

    try:
        proc = subprocess.run(cmd, env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return proc.returncode == 0 and os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 1024
    except Exception:
        return False


# --- Convert PDF to DOCX using pdf2docx (layout mode) ---
def _pdf_to_docx_layout(input_pdf, output_docx):
    try:
        from pdf2docx import Converter
        cv = Converter(input_pdf)
        cv.convert(output_docx, start=0, end=None)
        cv.close()
        return os.path.exists(output_docx) and os.path.getsize(output_docx) > 1024
    except Exception:
        return False


# --- Extract tables and append to DOCX (fallback) ---
def _append_tables(pdf_path, output_docx):
    try:
        tabula.environment_info.java_path = JAVA_EXE if os.path.exists(JAVA_EXE) else None
        tables = tabula.read_pdf(pdf_path, pages="all", multiple_tables=True, lattice=True, stream=False,
                                 java_options=["-Djava.awt.headless=true"], pandas_options={"dtype": str}) or []
        if not tables:
            tables = tabula.read_pdf(pdf_path, pages="all", multiple_tables=True, lattice=False, stream=True,
                                     java_options=["-Djava.awt.headless=true"], pandas_options={"dtype": str}) or []
    except Exception:
        tables = []

    if not tables:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    for tbl in page.extract_tables() or []:
                        tables.append(tbl)
        except Exception:
            pass

    if not tables:
        return

    doc = Document(output_docx)
    doc.add_page_break()
    doc.add_paragraph("Extracted Tables (auto)").bold = True

    for tbl in tables:
        try:
            rows = tbl.values.tolist() if hasattr(tbl, "values") else tbl
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    t.cell(i, j).text = "" if cell is None else str(cell)
            doc.add_paragraph()
        except Exception:
            continue

    doc.save(output_docx)


# === Main function called by app.py ===
def convert_pdf_to_docx(pdf_path, out_docx):
    pdf_path = str(Path(pdf_path))
    out_docx = str(Path(out_docx))

    if _pdf_has_text(pdf_path):
        ok = _pdf_to_docx_layout(pdf_path, out_docx)
        if ok:
            _append_tables(pdf_path, out_docx)
        return ok

    with tempfile.TemporaryDirectory() as td:
        searchable_pdf = os.path.join(td, "searchable.pdf")
        if not _ocr_to_searchable_pdf(pdf_path, searchable_pdf):
            return False
        ok = _pdf_to_docx_layout(searchable_pdf, out_docx)
        if ok:
            _append_tables(pdf_path, out_docx)
        return ok


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("pdf", help="Input PDF")
    ap.add_argument("docx", help="Output DOCX")
    args = ap.parse_args()
    print("SUCCESS" if convert_pdf_to_docx(args.pdf, args.docx) else "FAILED")
